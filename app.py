import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import re
import json
from io import StringIO
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# Import authentication module
from auth_module import auth_manager, require_auth, get_user_display_name, is_authenticated

# Document processing imports
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        USE_PDFPLUMBER = True
    except ImportError:
        PDF_AVAILABLE = False
        USE_PDFPLUMBER = False

# NLP imports
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.cluster import KMeans
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

try:
    from textblob import TextBlob
    TEXTBLOB_AVAILABLE = True
except ImportError:
    TEXTBLOB_AVAILABLE = False

try:
    import nltk
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

# Set style for better looking charts
plt.style.use('default')
sns.set_palette("husl")

# Streamlit App Configuration
st.set_page_config(
    page_title="AI Fit Scoring Dashboard",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        color: white;
    }
    .metric-card {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border-left: 4px solid #667eea;
    }
    .success-score {
        background: linear-gradient(45deg, #56ab2f, #a8e6cf);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
    }
    .warning-score {
        background: linear-gradient(45deg, #f7971e, #ffd200);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
    }
    .danger-score {
        background: linear-gradient(45deg, #e74c3c, #c0392b);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
    }
    .sidebar .sidebar-content {
        background: #f8f9fa;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
def init_session_state():
    """Initialize session state variables"""
    if 'scores' not in st.session_state:
        st.session_state.scores = {}
    if 'reasoning' not in st.session_state:
        st.session_state.reasoning = {}
    if 'nlp_analysis' not in st.session_state:
        st.session_state.nlp_analysis = {}
    if 'current_project_name' not in st.session_state:
        st.session_state.current_project_name = ""
    if 'analysis_history' not in st.session_state:
        st.session_state.analysis_history = []

init_session_state()

# Enhanced criteria with more sophisticated definitions
CRITERIA = {
    "Data Requirements": {
        "description": "Quality, quantity, and accessibility of data needed for AI implementation",
        "weight": 25,
        "detailed_description": "Evaluates whether sufficient, high-quality data exists or can be obtained for training AI models. Consider data volume, variety, velocity, and veracity.",
        "key_factors": ["Data availability", "Data quality", "Data labeling", "Data governance", "Privacy compliance"]
    },
    "Problem Complexity": {
        "description": "Suitability of the problem for AI/ML approaches and techniques",
        "weight": 20,
        "detailed_description": "Assesses if the problem has characteristics that make it suitable for AI solutions, such as pattern recognition, prediction, or optimization opportunities.",
        "key_factors": ["Pattern recognition potential", "Predictive elements", "Optimization opportunities", "Rule complexity", "Human expert availability"]
    },
    "Business Impact": {
        "description": "Expected ROI, strategic value, and business transformation potential",
        "weight": 20,
        "detailed_description": "Measures the potential business value, cost savings, revenue generation, and competitive advantage the AI solution could provide.",
        "key_factors": ["Revenue impact", "Cost reduction", "Efficiency gains", "Competitive advantage", "Scalability potential"]
    },
    "Technical Feasibility": {
        "description": "Technical requirements, infrastructure, and implementation complexity",
        "weight": 20,
        "detailed_description": "Evaluates technical constraints, required infrastructure, integration complexity, and availability of necessary AI/ML technologies and expertise.",
        "key_factors": ["Infrastructure readiness", "Integration complexity", "Technology maturity", "Team expertise", "Maintenance requirements"]
    },
    "Timeline Alignment": {
        "description": "Compatibility with business timelines and realistic development expectations",
        "weight": 15,
        "detailed_description": "Assesses whether the AI project timeline aligns with business needs, considering development complexity and realistic delivery expectations.",
        "key_factors": ["Development timeline", "Business urgency", "Iterative delivery", "Risk tolerance", "Change management"]
    }
}

class EnhancedNLPAnalyzer:
    """Enhanced NLP analyzer with improved algorithms and industry-specific insights"""
    
    def __init__(self):
        self.tfidf = None
        if SKLEARN_AVAILABLE:
            self.tfidf = TfidfVectorizer(
                max_features=200,
                stop_words='english',
                ngram_range=(1, 3),
                min_df=1,
                max_df=0.95
            )
        
        # Industry-specific AI readiness indicators
        self.ai_readiness_indicators = {
            'high_value': [
                'machine learning', 'artificial intelligence', 'deep learning', 'neural network',
                'predictive analytics', 'pattern recognition', 'computer vision', 'natural language',
                'recommendation system', 'anomaly detection', 'fraud detection', 'personalization',
                'optimization', 'automation', 'intelligent', 'smart system'
            ],
            'moderate_value': [
                'data analysis', 'statistics', 'algorithm', 'model', 'prediction', 'classification',
                'clustering', 'regression', 'decision tree', 'data mining', 'analytics',
                'dashboard', 'reporting', 'insights', 'trends'
            ],
            'low_value': [
                'manual process', 'simple calculation', 'basic reporting', 'static analysis',
                'rule-based', 'hardcoded', 'fixed logic', 'simple automation'
            ]
        }
        
        # Business impact keywords
        self.business_impact_keywords = {
            'revenue': ['revenue', 'sales', 'profit', 'income', 'monetization', 'pricing'],
            'cost_reduction': ['cost', 'expense', 'efficiency', 'optimize', 'streamline', 'reduce'],
            'customer': ['customer', 'user', 'client', 'satisfaction', 'experience', 'retention'],
            'competitive': ['competitive', 'advantage', 'market', 'differentiation', 'innovation'],
            'scale': ['scale', 'growth', 'expansion', 'scalable', 'volume', 'capacity']
        }
        
        # Technical feasibility indicators
        self.tech_feasibility_keywords = {
            'positive': [
                'existing infrastructure', 'proven technology', 'available data', 'skilled team',
                'cloud platform', 'api', 'integration', 'modular', 'scalable architecture'
            ],
            'negative': [
                'legacy system', 'limited budget', 'tight deadline', 'no data', 'complex integration',
                'regulatory constraints', 'security concerns', 'limited expertise'
            ]
        }
    
    def extract_keywords(self, text: str, top_k: int = 15) -> List[Tuple[str, float]]:
        """Extract important keywords using enhanced TF-IDF"""
        if not SKLEARN_AVAILABLE or not text:
            return []
        
        try:
            processed_text = self.preprocess_text(text)
            if len(processed_text.split()) < 5:
                return []
                
            tfidf_matrix = self.tfidf.fit_transform([processed_text])
            feature_names = self.tfidf.get_feature_names_out()
            tfidf_scores = tfidf_matrix.toarray()[0]
            
            # Get top keywords
            top_indices = np.argsort(tfidf_scores)[::-1][:top_k]
            keywords = [(feature_names[i], tfidf_scores[i]) for i in top_indices if tfidf_scores[i] > 0]
            
            return keywords
        except Exception as e:
            st.warning(f"Keyword extraction failed: {str(e)}")
            return []
    
    def analyze_sentiment(self, text: str) -> Optional[Dict]:
        """Enhanced sentiment analysis with confidence scoring"""
        if not TEXTBLOB_AVAILABLE or not text:
            return None
        
        try:
            # Process text in chunks to handle large documents
            chunks = [text[i:i+1000] for i in range(0, len(text), 1000)]
            polarities = []
            subjectivities = []
            
            for chunk in chunks[:5]:  # Limit to first 5 chunks
                blob = TextBlob(chunk)
                polarities.append(blob.sentiment.polarity)
                subjectivities.append(blob.sentiment.subjectivity)
            
            avg_polarity = np.mean(polarities)
            avg_subjectivity = np.mean(subjectivities)
            
            # Enhanced confidence calculation
            polarity_std = np.std(polarities)
            confidence_base = 0.5 + (0.5 - polarity_std)
            
            # Convert polarity to label with better thresholds
            if avg_polarity > 0.15:
                label = "POSITIVE"
                confidence = min(confidence_base + (avg_polarity * 0.3), 0.95)
            elif avg_polarity < -0.15:
                label = "NEGATIVE"
                confidence = min(confidence_base + (abs(avg_polarity) * 0.3), 0.95)
            else:
                label = "NEUTRAL"
                confidence = min(confidence_base, 0.85)
            
            return {
                'label': label,
                'confidence': confidence,
                'polarity': avg_polarity,
                'subjectivity': avg_subjectivity,
                'consistency': 1 - polarity_std  # How consistent sentiment is across text
            }
        except Exception as e:
            st.warning(f"Sentiment analysis failed: {str(e)}")
            return None
    
    def preprocess_text(self, text: str) -> str:
        """Enhanced text preprocessing"""
        if not text:
            return ""
        
        # Basic cleaning
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\-\.\,\;\:]', ' ', text)
        text = text.lower().strip()
        
        # Remove common noise
        noise_patterns = [
            r'\b\d{4}\b',  # Years
            r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b',  # Dates
            r'\b\w{1,2}\b'  # Very short words
        ]
        
        for pattern in noise_patterns:
            text = re.sub(pattern, ' ', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def calculate_ai_readiness_score(self, text: str) -> Dict[str, float]:
        """Calculate AI readiness based on keyword analysis"""
        text_lower = text.lower()
        
        scores = {}
        total_indicators = 0
        
        for category, keywords in self.ai_readiness_indicators.items():
            count = sum(1 for keyword in keywords if keyword in text_lower)
            scores[category] = count
            total_indicators += count
        
        # Calculate readiness percentage
        if total_indicators == 0:
            readiness_score = 0.3  # Default low score
        else:
            high_weight = scores.get('high_value', 0) * 3
            moderate_weight = scores.get('moderate_value', 0) * 2
            low_weight = scores.get('low_value', 0) * 0.5
            
            weighted_score = (high_weight + moderate_weight - low_weight) / (total_indicators + 1)
            readiness_score = min(max(weighted_score * 0.2, 0), 1)
        
        return {
            'score': readiness_score,
            'high_value_indicators': scores.get('high_value', 0),
            'moderate_value_indicators': scores.get('moderate_value', 0),
            'low_value_indicators': scores.get('low_value', 0),
            'total_indicators': total_indicators
        }
    
    def extract_business_metrics(self, text: str) -> Dict[str, List[str]]:
        """Extract business impact indicators and metrics"""
        findings = {category: [] for category in self.business_impact_keywords.keys()}
        text_lower = text.lower()
        
        for category, keywords in self.business_impact_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    # Find context around the keyword
                    pattern = rf'.{{0,50}}{keyword}.{{0,50}}'
                    matches = re.findall(pattern, text_lower)
                    findings[category].extend(matches[:2])  # Limit to 2 matches per keyword
        
        return findings
    
    def assess_technical_feasibility(self, text: str) -> Dict[str, float]:
        """Assess technical feasibility based on text analysis"""
        text_lower = text.lower()
        
        positive_score = 0
        negative_score = 0
        
        for keyword in self.tech_feasibility_keywords['positive']:
            if keyword in text_lower:
                positive_score += 1
        
        for keyword in self.tech_feasibility_keywords['negative']:
            if keyword in text_lower:
                negative_score += 1
        
        # Calculate feasibility score (0-1)
        total_mentions = positive_score + negative_score
        if total_mentions == 0:
            feasibility_score = 0.5  # Neutral when no indicators
        else:
            feasibility_score = positive_score / total_mentions
        
        return {
            'score': feasibility_score,
            'positive_indicators': positive_score,
            'negative_indicators': negative_score,
            'confidence': min(total_mentions / 5, 1.0)  # Higher confidence with more indicators
        }
    
    def comprehensive_analysis(self, text: str) -> Dict:
        """Run comprehensive analysis combining all NLP features"""
        analysis = {
            'text_stats': {
                'word_count': len(text.split()),
                'char_count': len(text),
                'avg_sentence_length': len(text.split()) / max(text.count('.') + text.count('!') + text.count('?'), 1)
            },
            'keywords': self.extract_keywords(text),
            'sentiment': self.analyze_sentiment(text),
            'ai_readiness': self.calculate_ai_readiness_score(text),
            'business_metrics': self.extract_business_metrics(text),
            'technical_feasibility': self.assess_technical_feasibility(text),
            'project_type': self.detect_project_type(text),
            'complexity_indicators': self.detect_complexity_indicators(text)
        }
        
        return analysis
    
    def detect_project_type(self, text: str) -> Dict[str, float]:
        """Enhanced project type detection with confidence scores"""
        text_lower = text.lower()
        
        project_patterns = {
            'Machine Learning/AI': {
                'keywords': ['machine learning', 'ai', 'artificial intelligence', 'neural network', 'deep learning', 'model training', 'prediction', 'classification', 'recommendation'],
                'weight': 1.0
            },
            'Data Analytics': {
                'keywords': ['analytics', 'dashboard', 'reporting', 'visualization', 'business intelligence', 'data analysis', 'insights'],
                'weight': 0.8
            },
            'Web Application': {
                'keywords': ['web', 'website', 'frontend', 'backend', 'api', 'browser', 'responsive'],
                'weight': 0.7
            },
            'Mobile Application': {
                'keywords': ['mobile', 'app', 'ios', 'android', 'smartphone', 'tablet'],
                'weight': 0.7
            },
            'Automation': {
                'keywords': ['automation', 'workflow', 'process', 'streamline', 'optimize', 'robotic'],
                'weight': 0.6
            },
            'E-commerce': {
                'keywords': ['ecommerce', 'shop', 'store', 'payment', 'cart', 'inventory', 'marketplace'],
                'weight': 0.5
            }
        }
        
        scores = {}
        for project_type, config in project_patterns.items():
            score = 0
            for keyword in config['keywords']:
                if keyword in text_lower:
                    score += config['weight']
            scores[project_type] = score
        
        # Normalize scores
        max_score = max(scores.values()) if scores.values() else 1
        normalized_scores = {k: v/max_score for k, v in scores.items()}
        
        return normalized_scores
    
    def detect_complexity_indicators(self, text: str) -> Dict[str, int]:
        """Detect complexity indicators in the project description"""
        text_lower = text.lower()
        
        complexity_patterns = {
            'integration_points': len(re.findall(r'integrat\w+|connect\w+|api\w*|interface\w*', text_lower)),
            'data_sources': len(re.findall(r'database\w*|data\s+source\w*|feed\w*|stream\w*', text_lower)),
            'user_types': len(re.findall(r'user\w*|role\w*|permission\w*|access\w*', text_lower)),
            'business_rules': len(re.findall(r'rule\w*|condition\w*|logic\w*|workflow\w*', text_lower)),
            'external_deps': len(re.findall(r'third.party|external|vendor|partner', text_lower))
        }
        
        return complexity_patterns

def extract_text_from_file(uploaded_file) -> Tuple[bool, str, str]:
    """Enhanced file processing with better error handling"""
    file_extension = uploaded_file.name.lower().split('.')[-1]
    
    try:
        if file_extension in ['txt', 'md']:
            text = str(uploaded_file.read(), "utf-8")
            return True, text, ""
            
        elif file_extension == 'docx':
            if not DOCX_AVAILABLE:
                return False, "", "python-docx library not available. Please install it."
            
            doc = Document(uploaded_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            
            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += "\n" + " | ".join(row_text)
            
            return True, text, ""
            
        elif file_extension == 'pdf':
            if not PDF_AVAILABLE:
                return False, "", "PDF processing library not available. Please install PyMuPDF or pdfplumber."
            
            text = ""
            try:
                if 'USE_PDFPLUMBER' in globals() and USE_PDFPLUMBER:
                    import pdfplumber
                    with pdfplumber.open(uploaded_file) as pdf:
                        for page_num, page in enumerate(pdf.pages[:10]):  # Limit to first 10 pages
                            page_text = page.extract_text()
                            if page_text:
                                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                else:
                    pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                    for page_num in range(min(10, pdf_document.page_count)):
                        page = pdf_document[page_num]
                        text += f"\n--- Page {page_num + 1} ---\n{page.get_text()}"
                    pdf_document.close()
                
                return True, text, ""
            except Exception as e:
                return False, "", f"PDF processing error: {str(e)}"
            
        else:
            return False, "", f"Unsupported file format: {file_extension}. Supported: PDF, DOCX, TXT, MD"
            
    except Exception as e:
        return False, "", f"Error processing {file_extension.upper()} file: {str(e)}"

def advanced_ai_scoring(text: str, criteria_name: str, nlp_analyzer: EnhancedNLPAnalyzer) -> Tuple[int, str, Dict]:
    """Enhanced AI-powered scoring with detailed analysis"""
    if not text:
        return 3, "No text provided for analysis", {}
    
    # Get comprehensive analysis
    analysis = nlp_analyzer.comprehensive_analysis(text)
    
    # Criteria-specific scoring logic
    if criteria_name == "Data Requirements":
        score, reasoning = score_data_requirements(text, analysis)
    elif criteria_name == "Problem Complexity":
        score, reasoning = score_problem_complexity(text, analysis)
    elif criteria_name == "Business Impact":
        score, reasoning = score_business_impact(text, analysis)
    elif criteria_name == "Technical Feasibility":
        score, reasoning = score_technical_feasibility(text, analysis)
    elif criteria_name == "Timeline Alignment":
        score, reasoning = score_timeline_alignment(text, analysis)
    else:
        score, reasoning = 3, "Unknown criteria"
    
    return score, reasoning, analysis

def score_data_requirements(text: str, analysis: Dict) -> Tuple[int, str]:
    """Score data requirements based on NLP analysis"""
    text_lower = text.lower()
    base_score = 3
    reasoning_parts = []
    
    # Data availability indicators
    data_positive = ['existing data', 'historical data', 'data warehouse', 'data lake', 'structured data', 'clean data']
    data_negative = ['no data', 'limited data', 'poor quality', 'unstructured', 'missing data']
    
    positive_count = sum(1 for term in data_positive if term in text_lower)
    negative_count = sum(1 for term in data_negative if term in text_lower)
    
    # AI readiness score influence
    ai_readiness = analysis.get('ai_readiness', {})
    readiness_score = ai_readiness.get('score', 0.3)
    
    # Calculate final score
    data_score_adjustment = (positive_count * 0.5) - (negative_count * 0.4) + (readiness_score * 2)
    final_score = max(1, min(5, int(base_score + data_score_adjustment)))
    
    # Build reasoning
    if positive_count > 0:
        reasoning_parts.append(f"✓ {positive_count} positive data indicators found")
    if negative_count > 0:
        reasoning_parts.append(f"⚠ {negative_count} data concerns identified")
    if readiness_score > 0.6:
        reasoning_parts.append(f"✓ High AI readiness score ({readiness_score:.2f})")
    
    reasoning = "; ".join(reasoning_parts) if reasoning_parts else "Limited data-related content detected"
    
    return final_score, reasoning

def score_problem_complexity(text: str, analysis: Dict) -> Tuple[int, str]:
    """Score problem complexity for AI suitability"""
    complexity_indicators = analysis.get('complexity_indicators', {})
    ai_readiness = analysis.get('ai_readiness', {})
    
    # Base score calculation
    complexity_score = sum(min(v, 3) for v in complexity_indicators.values()) / 5
    ai_score = ai_readiness.get('score', 0.3) * 3
    
    final_score = max(1, min(5, int(2 + complexity_score + ai_score)))
    
    reasoning_parts = []
    if complexity_indicators.get('integration_points', 0) > 2:
        reasoning_parts.append("✓ Multiple integration points detected")
    if ai_readiness.get('high_value_indicators', 0) > 0:
        reasoning_parts.append(f"✓ {ai_readiness['high_value_indicators']} AI-suitable patterns found")
    
    reasoning = "; ".join(reasoning_parts) if reasoning_parts else "Moderate complexity detected"
    
    return final_score, reasoning

def score_business_impact(text: str, analysis: Dict) -> Tuple[int, str]:
    """Score business impact potential"""
    business_metrics = analysis.get('business_metrics', {})
    sentiment = analysis.get('sentiment', {})
    
    # Count business impact indicators
    impact_score = 0
    for category, findings in business_metrics.items():
        impact_score += min(len(findings), 2) * 0.3
    
    # Sentiment influence
    if sentiment and sentiment.get('label') == 'POSITIVE':
        impact_score += sentiment.get('confidence', 0) * 0.5
    
    final_score = max(1, min(5, int(2.5 + impact_score)))
    
    reasoning_parts = []
    high_impact_categories = [cat for cat, findings in business_metrics.items() if len(findings) > 0]
    if high_impact_categories:
        reasoning_parts.append(f"✓ Business impact in: {', '.join(high_impact_categories[:3])}")
    
    reasoning = "; ".join(reasoning_parts) if reasoning_parts else "Limited business impact indicators"
    
    return final_score, reasoning

def score_technical_feasibility(text: str, analysis: Dict) -> Tuple[int, str]:
    """Score technical feasibility"""
    tech_feasibility = analysis.get('technical_feasibility', {})
    feasibility_score = tech_feasibility.get('score', 0.5)
    confidence = tech_feasibility.get('confidence', 0)
    
    final_score = max(1, min(5, int(1 + feasibility_score * 4)))
    
    reasoning_parts = []
    if tech_feasibility.get('positive_indicators', 0) > 0:
        reasoning_parts.append(f"✓ {tech_feasibility['positive_indicators']} positive tech indicators")
    if tech_feasibility.get('negative_indicators', 0) > 0:
        reasoning_parts.append(f"⚠ {tech_feasibility['negative_indicators']} technical concerns")
    if confidence > 0.7:
        reasoning_parts.append("✓ High confidence assessment")
    
    reasoning = "; ".join(reasoning_parts) if reasoning_parts else "Standard technical feasibility"
    
    return final_score, reasoning

def score_timeline_alignment(text: str, analysis: Dict) -> Tuple[int, str]:
    """Score timeline alignment"""
    text_lower = text.lower()
    
    # Timeline indicators
    urgent_terms = ['urgent', 'asap', 'immediately', 'rush', 'critical']
    reasonable_terms = ['phased', 'iterative', 'gradual', 'milestone', 'sprint']
    
    urgent_count = sum(1 for term in urgent_terms if term in text_lower)
    reasonable_count = sum(1 for term in reasonable_terms if term in text_lower)
    
    # Calculate score
    if urgent_count > reasonable_count:
        final_score = max(1, 3 - urgent_count)
    else:
        final_score = min(5, 3 + reasonable_count)
    
    reasoning_parts = []
    if urgent_count > 0:
        reasoning_parts.append(f"⚠ {urgent_count} urgency indicators")
    if reasonable_count > 0:
        reasoning_parts.append(f"✓ {reasonable_count} structured timeline indicators")
    
    reasoning = "; ".join(reasoning_parts) if reasoning_parts else "Standard timeline expectations"
    
    return final_score, reasoning

def render_authentication_ui():
    """Render authentication interface"""
    st.sidebar.title("🔐 Authentication")
    
    if is_authenticated():
        user_name = get_user_display_name()
        st.sidebar.success(f"Welcome, {user_name}!")
        
        if st.sidebar.button("🚪 Sign Out", key="signout"):
            if auth_manager.sign_out():
                st.rerun()
        
        return True
    else:
        auth_tab1, auth_tab2 = st.sidebar.tabs(["🔑 Sign In", "📝 Sign Up"])
        
        with auth_tab1:
            with st.form("signin_form"):
                email = st.text_input("Email")
                password = st.text_input("Password", type="password")
                
                if st.form_submit_button("🔑 Sign In"):
                    success, message = auth_manager.sign_in(email, password)
                    if success:
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)
        
        with auth_tab2:
            with st.form("signup_form"):
                full_name = st.text_input("Full Name")
                email = st.text_input("Email")
                password = st.text_input("Password", type="password")
                confirm_password = st.text_input("Confirm Password", type="password")
                
                if st.form_submit_button("📝 Create Account"):
                    if password != confirm_password:
                        st.error("Passwords do not match")
                    elif len(password) < 6:
                        st.error("Password must be at least 6 characters")
                    else:
                        success, message = auth_manager.sign_up(email, password, full_name)
                        if success:
                            st.success(message)
                        else:
                            st.error(message)
        
        return False

def render_user_dashboard():
    """Render the user dashboard with project history"""
    st.header(f"👋 Welcome back, {get_user_display_name()}!")
    
    # Get user projects
    projects = auth_manager.get_user_projects(limit=20)
    
    if projects:
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📊 Total Projects", len(projects))
        with col2:
            recent_projects = [p for p in projects if 
                             (datetime.now() - datetime.fromisoformat(p['created_at'].replace('Z', '+00:00').replace('+00:00', ''))).days <= 7]
            st.metric("📅 This Week", len(recent_projects))
        with col3:
            avg_score = np.mean([p.get('overall_score', 0) for p in projects if p.get('overall_score')])
            st.metric("⭐ Avg Score", f"{avg_score:.1f}%" if avg_score else "N/A")
        
        st.subheader("📁 Recent Projects")
        
        for project in projects[:5]:
            with st.expander(f"🎯 {project.get('project_name', 'Unnamed Project')} - {project.get('overall_score', 0):.0f}%"):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"**Type:** {project.get('project_type', 'Unknown')}")
                    st.write(f"**Created:** {project.get('created_at', '')[:10]}")
                    if project.get('description'):
                        st.write(f"**Description:** {project['description'][:100]}...")
                with col2:
                    if st.button("🗑️ Delete", key=f"delete_{project['id']}"):
                        success, message = auth_manager.delete_project(project['id'])
                        if success:
                            st.success(message)
                            st.rerun()
                        else:
                            st.error(message)
    else:
        st.info("📈 No projects yet. Create your first AI fit assessment!")
        st.markdown("**Get started by:**")
        st.markdown("- 📄 Uploading a project document")
        st.markdown("- ✍️ Using manual scoring")
        st.markdown("- 🧠 Trying our advanced NLP analysis")

@require_auth
def save_project_to_db(project_data: Dict, scoring_data: Dict, nlp_analysis: Dict = None):
    """Save project and scoring results to database"""
    # Save project
    success, message, project_id = auth_manager.save_project(project_data)
    
    if success and project_id:
        # Save scoring results
        scoring_success, scoring_message = auth_manager.save_scoring_results(
            project_id, scoring_data, nlp_analysis
        )
        
        if scoring_success:
            st.success(f"✅ Project saved successfully! {message}")
        else:
            st.warning(f"Project saved but scoring failed: {scoring_message}")
    else:
        st.error(f"Failed to save project: {message}")

def main():
    """Main application function"""
    
    # Authentication check
    if not render_authentication_ui():
        st.markdown("""
        <div class="main-header">
            <h1>🎯 AI Fit Scoring Platform</h1>
            <p>Intelligent assessment tool to determine if AI is the right fit for your project</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        ### 🚀 Features
        - 🧠 **Advanced NLP Analysis** - Automated scoring using machine learning
        - 📊 **Comprehensive Scoring** - 5-criteria evaluation framework
        - 📈 **Visual Analytics** - Charts and insights for better decision making
        - 💾 **Project History** - Save and track your assessments
        - 📄 **Document Processing** - Support for PDF, DOCX, TXT, MD files
        
        ### 🔐 Sign up now to get started!
        """)
        return
    
    # Main app logic for authenticated users
    st.markdown("""
    <div class="main-header">
        <h1>🎯 AI Fit Scoring Dashboard</h1>
        <p>Advanced NLP-powered assessment platform</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Navigation tabs
    tab1, tab2, tab3, tab4 = st.tabs(["🏠 Dashboard", "📄 Analyze Document", "✍️ Manual Scoring", "⚙️ Settings"])
    
    with tab1:
        render_user_dashboard()
    
    with tab2:
        render_document_analysis()
    
    with tab3:
        render_manual_scoring()
    
    with tab4:
        render_settings()

def render_document_analysis():
    """Document upload and analysis interface"""
    st.subheader("📤 Document Upload & Advanced NLP Analysis")
    
    # Project name input
    project_name = st.text_input("🏷️ Project Name", 
                                placeholder="Enter a name for this project...",
                                key="doc_project_name")
    
    uploaded_file = st.file_uploader(
        "Upload your project document",
        type=['pdf', 'docx', 'txt', 'md'],
        help="Supported formats: PDF, DOCX, TXT, MD",
        key="document_uploader"
    )
    
    if uploaded_file is not None:
        # File info display
        col1, col2, col3 = st.columns(3)
        with col1:
            st.success(f"✅ {uploaded_file.name}")
        with col2:
            st.info(f"📊 {uploaded_file.size / 1024:.1f} KB")
        
        # Extract text
        with st.spinner("🔄 Processing document..."):
            success, extracted_text, error_message = extract_text_from_file(uploaded_file)
        
        if success and extracted_text.strip():
            # Initialize enhanced NLP analyzer
            nlp_analyzer = EnhancedNLPAnalyzer()
            
            # Document statistics
            word_count = len(extracted_text.split())
            char_count = len(extracted_text)
            
            with col3:
                st.metric("📝 Words", f"{word_count:,}")
            
            # Run comprehensive analysis
            with st.spinner("🧠 Running comprehensive NLP analysis..."):
                comprehensive_analysis = nlp_analyzer.comprehensive_analysis(extracted_text)
            
            # Display key metrics
            col4, col5, col6 = st.columns(3)
            with col4:
                st.metric("🔤 Characters", f"{char_count:,}")
            with col5:
                project_types = comprehensive_analysis.get('project_type', {})
                top_type = max(project_types, key=project_types.get) if project_types else "Unknown"
                st.metric("🎯 Project Type", top_type)
            with col6:
                ai_readiness = comprehensive_analysis.get('ai_readiness', {})
                readiness_score = ai_readiness.get('score', 0) * 100
                st.metric("🤖 AI Readiness", f"{readiness_score:.0f}%")
            
            # Advanced NLP Analysis Results
            with st.expander("🧠 Comprehensive Analysis Results", expanded=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    # Keywords and sentiment
                    keywords = comprehensive_analysis.get('keywords', [])
                    if keywords:
                        st.subheader("🔍 Key Terms (TF-IDF)")
                        for keyword, score in keywords[:8]:
                            st.write(f"• **{keyword}** ({score:.3f})")
                    
                    # AI Readiness breakdown
                    ai_readiness = comprehensive_analysis.get('ai_readiness', {})
                    if ai_readiness:
                        st.subheader("🤖 AI Readiness Analysis")
                        st.write(f"**High-value indicators:** {ai_readiness.get('high_value_indicators', 0)}")
                        st.write(f"**Moderate indicators:** {ai_readiness.get('moderate_value_indicators', 0)}")
                        st.write(f"**Low-value indicators:** {ai_readiness.get('low_value_indicators', 0)}")
                
                with col2:
                    # Sentiment and business metrics
                    sentiment = comprehensive_analysis.get('sentiment', {})
                    if sentiment:
                        st.subheader("😊 Sentiment Analysis")
                        st.write(f"**Label:** {sentiment['label']}")
                        st.write(f"**Confidence:** {sentiment['confidence']:.2f}")
                        st.write(f"**Consistency:** {sentiment.get('consistency', 0):.2f}")
                    
                    # Technical feasibility
                    tech_feasibility = comprehensive_analysis.get('technical_feasibility', {})
                    if tech_feasibility:
                        st.subheader("⚙️ Technical Feasibility")
                        st.write(f"**Score:** {tech_feasibility.get('score', 0):.2f}")
                        st.write(f"**Positive indicators:** {tech_feasibility.get('positive_indicators', 0)}")
                        st.write(f"**Concerns:** {tech_feasibility.get('negative_indicators', 0)}")
            
            # Document preview
            with st.expander("📖 Document Preview"):
                preview_text = extracted_text[:1500] + "..." if len(extracted_text) > 1500 else extracted_text
                st.text_area("Content Preview:", preview_text, height=200, disabled=True)
            
            # Analysis and scoring buttons
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("🧠 Advanced AI Scoring", type="primary", key="ai_score"):
                    if not project_name.strip():
                        st.error("Please enter a project name")
                    else:
                        with st.spinner("🔍 Running advanced AI analysis..."):
                            # Run enhanced scoring for each criteria
                            for criteria_name in CRITERIA.keys():
                                score, reasoning, analysis = advanced_ai_scoring(
                                    extracted_text, criteria_name, nlp_analyzer
                                )
                                st.session_state.scores[criteria_name] = score
                                st.session_state.reasoning[criteria_name] = reasoning
                            
                            # Store comprehensive analysis
                            st.session_state.nlp_analysis = comprehensive_analysis
                            st.session_state.current_project_name = project_name
                        
                        st.success("🎯 Advanced AI scoring completed!")
                        st.rerun()
            
            with col2:
                if st.button("💾 Save Project", key="save_project"):
                    if not project_name.strip():
                        st.error("Please enter a project name")
                    elif not st.session_state.scores:
                        st.error("Please run scoring analysis first")
                    else:
                        # Calculate overall score
                        total_weighted_score = calculate_overall_score(st.session_state.scores)
                        
                        # Prepare project data
                        project_data = {
                            'project_name': project_name,
                            'project_type': top_type,
                            'description': extracted_text[:500],
                            'overall_score': total_weighted_score,
                            'document_name': uploaded_file.name,
                            'word_count': word_count
                        }
                        
                        # Prepare scoring data
                        scoring_data = {
                            'overall_score': total_weighted_score,
                            'scoring_method': 'advanced_nlp'
                        }
                        
                        # Add individual criteria scores
                        for criteria_name in CRITERIA.keys():
                            if criteria_name in st.session_state.scores:
                                key_base = criteria_name.lower().replace(' ', '_')
                                scoring_data[f'{key_base}_score'] = st.session_state.scores[criteria_name]
                                scoring_data[f'{key_base}_reasoning'] = st.session_state.reasoning.get(criteria_name, '')
                        
                        # Save to database
                        save_project_to_db(project_data, scoring_data, comprehensive_analysis)
            
            with col3:
                if st.button("🔄 Reset Analysis", key="reset"):
                    st.session_state.scores = {}
                    st.session_state.reasoning = {}
                    st.session_state.nlp_analysis = {}
                    st.session_state.current_project_name = ""
                    st.rerun()
                    
        else:
            st.error(f"❌ Error: {error_message}")

def render_manual_scoring():
    """Manual scoring interface"""
    st.subheader("✍️ Manual Project Scoring")
    
    # Project name input
    project_name = st.text_input("🏷️ Project Name", 
                                placeholder="Enter a name for this project...",
                                key="manual_project_name")
    
    # Project description
    project_description = st.text_area("📝 Project Description (Optional)",
                                      placeholder="Describe your project to help with analysis...",
                                      height=100,
                                      key="manual_description")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.write("**📊 Scoring Interface**")
        
        for criteria_name, criteria_info in CRITERIA.items():
            with st.expander(f"📋 {criteria_name} ({criteria_info['weight']}%)", expanded=True):
                st.write(f"**Description:** {criteria_info['description']}")
                st.write(f"**Detailed:** {criteria_info['detailed_description']}")
                
                # Key factors
                st.write("**Key Factors:**")
                for factor in criteria_info['key_factors']:
                    st.write(f"• {factor}")
                
                # Get current values
                current_score = st.session_state.scores.get(criteria_name, 3)
                current_reasoning = st.session_state.reasoning.get(criteria_name, "")
                
                score = st.slider(
                    f"Score for {criteria_name}",
                    1, 5, current_score,
                    format="%d",
                    key=f"manual_score_{criteria_name}",
                    help="1=Poor, 2=Below Average, 3=Average, 4=Good, 5=Excellent"
                )
                
                reasoning = st.text_area(
                    f"Reasoning for {criteria_name}",
                    current_reasoning,
                    height=80,
                    key=f"manual_reasoning_{criteria_name}",
                    placeholder="Explain your scoring rationale..."
                )
                
                # Update session state
                st.session_state.scores[criteria_name] = score
                st.session_state.reasoning[criteria_name] = reasoning
    
    with col2:
        st.write("**📈 Live Results**")
        if st.session_state.scores:
            total_score = calculate_overall_score(st.session_state.scores)
            
            # Score display
            if total_score >= 70:
                st.success(f"## {total_score:.0f}%\n**Excellent AI Fit** 🟢")
            elif total_score >= 60:
                st.info(f"## {total_score:.0f}%\n**Good AI Fit** 🔵")
            elif total_score >= 40:
                st.warning(f"## {total_score:.0f}%\n**Moderate AI Fit** 🟡")
            else:
                st.error(f"## {total_score:.0f}%\n**Low AI Fit** 🔴")
            
            # Individual scores
            st.write("**Individual Scores:**")
            for criteria_name, criteria_info in CRITERIA.items():
                if criteria_name in st.session_state.scores:
                    score = st.session_state.scores[criteria_name]
                    percentage = (score - 1) / 4 * 100
                    st.write(f"• {criteria_name}: {percentage:.0f}%")
        
        # Action buttons
        if st.button("💾 Save Manual Project", key="save_manual", type="primary"):
            if not project_name.strip():
                st.error("Please enter a project name")
            elif not st.session_state.scores:
                st.error("Please complete the scoring")
            else:
                # Calculate overall score
                total_weighted_score = calculate_overall_score(st.session_state.scores)
                
                # Prepare project data
                project_data = {
                    'project_name': project_name,
                    'project_type': 'Manual Assessment',
                    'description': project_description or 'Manual scoring assessment',
                    'overall_score': total_weighted_score
                }
                
                # Prepare scoring data
                scoring_data = {
                    'overall_score': total_weighted_score,
                    'scoring_method': 'manual'
                }
                
                # Add individual criteria scores
                for criteria_name in CRITERIA.keys():
                    if criteria_name in st.session_state.scores:
                        key_base = criteria_name.lower().replace(' ', '_')
                        scoring_data[f'{key_base}_score'] = st.session_state.scores[criteria_name]
                        scoring_data[f'{key_base}_reasoning'] = st.session_state.reasoning.get(criteria_name, '')
                
                # Save to database
                save_project_to_db(project_data, scoring_data)

def render_settings():
    """Settings and profile management"""
    st.subheader("⚙️ Settings & Profile")
    
    # User profile section
    st.write("**👤 User Profile**")
    profile = auth_manager.get_user_profile()
    current_user = auth_manager.get_current_user()
    
    if current_user:
        with st.form("profile_form"):
            full_name = st.text_input("Full Name", 
                                     value=profile.get('full_name', '') if profile else '')
            company = st.text_input("Company", 
                                   value=profile.get('company', '') if profile else '')
            role = st.text_input("Role/Title", 
                                value=profile.get('role', '') if profile else '')
            
            if st.form_submit_button("💾 Update Profile"):
                profile_data = {
                    'full_name': full_name,
                    'company': company,
                    'role': role
                }
                
                success, message = auth_manager.update_user_profile(profile_data)
                if success:
                    st.success(message)
                    st.rerun()
                else:
                    st.error(message)
    
    st.markdown("---")
    
    # Feedback section
    st.write("**💬 Feedback**")
    with st.form("feedback_form"):
        feedback_type = st.selectbox("Feedback Type", 
                                    ["General", "Bug Report", "Feature Request", "Improvement"])
        rating = st.slider("Overall Rating", 1, 5, 4)
        feedback_text = st.text_area("Your Feedback", 
                                    placeholder="Tell us what you think...")
        
        if st.form_submit_button("📤 Submit Feedback"):
            if feedback_text:
                feedback_data = {
                    'feedback_type': feedback_type.lower().replace(' ', '_'),
                    'rating': rating,
                    'feedback_text': feedback_text
                }
                
                success, message = auth_manager.submit_feedback(feedback_data)
                if success:
                    st.success(message)
                else:
                    st.error(message)
            else:
                st.error("Please enter your feedback")
    
    st.markdown("---")
    
    # System status
    st.write("**🔧 System Status**")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.write(f"📄 **PDF Processing:** {'✅' if PDF_AVAILABLE else '❌'}")
        st.write(f"📄 **DOCX Processing:** {'✅' if DOCX_AVAILABLE else '❌'}")
    with col2:
        st.write(f"🔍 **TF-IDF Analysis:** {'✅' if SKLEARN_AVAILABLE else '❌'}")
        st.write(f"😊 **Sentiment Analysis:** {'✅' if TEXTBLOB_AVAILABLE else '❌'}")
    with col3:
        st.write(f"🧠 **NLP Processing:** {'✅' if SKLEARN_AVAILABLE and TEXTBLOB_AVAILABLE else '❌'}")
        st.write(f"🔐 **Authentication:** {'✅' if auth_manager.is_initialized() else '❌'}")

def calculate_overall_score(scores: Dict[str, int]) -> float:
    """Calculate weighted overall score"""
    total_weighted_score = 0
    for criteria_name, criteria_info in CRITERIA.items():
        if criteria_name in scores:
            score = scores[criteria_name]
            percentage = (score - 1) / 4 * 100
            weight = criteria_info["weight"]
            total_weighted_score += percentage * (weight / 100)
    
    return total_weighted_score

# Results display function
def display_results():
    """Display comprehensive results section"""
    if not st.session_state.scores:
        return
    
    st.markdown("---")
    st.subheader("📊 AI Fit Assessment Results")
    
    # Calculate overall score
    total_weighted_score = calculate_overall_score(st.session_state.scores)
    
    # Display main score with enhanced styling
    score_col1, score_col2, score_col3 = st.columns([1, 2, 1])
    with score_col2:
        if total_weighted_score >= 70:
            st.markdown(f"""
            <div class="success-score">
                <h2>🟢 {total_weighted_score:.0f}% - Excellent AI Fit</h2>
                <p>✅ <strong>Highly Recommended</strong> - This project shows strong potential for AI implementation.</p>
            </div>
            """, unsafe_allow_html=True)
        elif total_weighted_score >= 60:
            st.markdown(f"""
            <div class="success-score">
                <h2>🔵 {total_weighted_score:.0f}% - Good AI Fit</h2>
                <p>✅ <strong>Recommended</strong> - This project has good AI potential with proper planning.</p>
            </div>
            """, unsafe_allow_html=True)
        elif total_weighted_score >= 40:
            st.markdown(f"""
            <div class="warning-score">
                <h2>🟡 {total_weighted_score:.0f}% - Moderate AI Fit</h2>
                <p>⚠️ <strong>Proceed with Caution</strong> - Evaluate risks and benefits carefully.</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="danger-score">
                <h2>🔴 {total_weighted_score:.0f}% - Low AI Fit</h2>
                <p>❌ <strong>Not Recommended</strong> - Consider alternative approaches.</p>
            </div>
            """, unsafe_allow_html=True)
    
    # NLP Insights Panel
    if st.session_state.nlp_analysis:
        st.subheader("🧠 NLP Analysis Summary")
        nlp_data = st.session_state.nlp_analysis
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            sentiment = nlp_data.get('sentiment', {})
            if sentiment:
                st.metric("Sentiment", sentiment.get('label', 'Unknown'))
        with col2:
            keywords = nlp_data.get('keywords', [])
            st.metric("Key Terms", len(keywords))
        with col3:
            ai_readiness = nlp_data.get('ai_readiness', {})
            readiness_score = ai_readiness.get('score', 0) * 100
            st.metric("AI Readiness", f"{readiness_score:.0f}%")
        with col4:
            project_types = nlp_data.get('project_type', {})
            top_type = max(project_types, key=project_types.get) if project_types else "Unknown"
            st.metric("Project Type", top_type)
    
    # Detailed scoring breakdown
    render_detailed_breakdown()
    
    # Visualization
    render_visualizations()
    
    # Export functionality
    render_export_section()

def render_detailed_breakdown():
    """Render detailed score breakdown"""
    st.subheader("📋 Detailed Score Breakdown")
    
    for criteria_name, criteria_info in CRITERIA.items():
        if criteria_name in st.session_state.scores:
            score = st.session_state.scores[criteria_name]
            percentage = (score - 1) / 4 * 100
            reasoning = st.session_state.reasoning.get(criteria_name, "No reasoning provided")
            
            # Color coding
            if percentage >= 75:
                color = "🟢"
                status = "Excellent"
            elif percentage >= 50:
                color = "🔵" 
                status = "Good"
            elif percentage >= 25:
                color = "🟡"
                status = "Average"
            else:
                color = "🔴"
                status = "Poor"
            
            with st.container():
                st.markdown(f"""
                <div class="metric-card">
                    <h4>{color} {criteria_name} - {percentage:.0f}% ({status})</h4>
                    <p><em>Weight: {criteria_info['weight']}% | Raw Score: {score}/5</em></p>
                    <p><strong>Reasoning:</strong> {reasoning}</p>
                </div>
                """, unsafe_allow_html=True)

def render_visualizations():
    """Render charts and visualizations"""
    st.markdown("---")
    st.subheader("📊 Visual Analysis")
    
    # Prepare data
    criteria_names = list(st.session_state.scores.keys())
    scores = list(st.session_state.scores.values())
    percentages = [(score - 1) / 4 * 100 for score in scores]
    
    # Create charts
    chart_col1, chart_col2 = st.columns(2)
    
    with chart_col1:
        # Enhanced Bar Chart
        fig, ax = plt.subplots(figsize=(12, 8))
        colors = ['#d32f2f' if p < 40 else '#f57c00' if p < 60 else '#388e3c' if p < 75 else '#1976d2' for p in percentages]
        
        bars = ax.bar(range(len(criteria_names)), percentages, color=colors, alpha=0.8, edgecolor='white', linewidth=2)
        ax.set_xlabel('Criteria', fontsize=12, fontweight='bold')
        ax.set_ylabel('Score (%)', fontsize=12, fontweight='bold')
        ax.set_title('AI Fit Score Breakdown', fontsize=16, fontweight='bold', pad=20)
        ax.set_xticks(range(len(criteria_names)))
        ax.set_xticklabels([name.replace(' ', '\n') for name in criteria_names], fontsize=10)
        ax.set_ylim(0, 100)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        
        # Add score labels on bars
        for bar, percentage in zip(bars, percentages):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 2,
                   f'{percentage:.0f}%', ha='center', va='bottom', fontweight='bold', fontsize=11)
        
        # Add horizontal reference lines
        ax.axhline(y=40, color='red', linestyle='--', alpha=0.7, label='Poor Threshold')
        ax.axhline(y=60, color='orange', linestyle='--', alpha=0.7, label='Good Threshold')
        ax.axhline(y=75, color='green', linestyle='--', alpha=0.7, label='Excellent Threshold')
        ax.legend(loc='upper right')
        
        plt.tight_layout()
        st.pyplot(fig)
    
    with chart_col2:
        # Enhanced Radar Chart
        fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(projection='polar'))
        
        angles = np.linspace(0, 2 * np.pi, len(criteria_names), endpoint=False).tolist()
        percentages_plot = percentages + [percentages[0]]
        angles += angles[:1]
        
        ax.plot(angles, percentages_plot, 'o-', linewidth=3, color='#1976d2', markersize=8)
        ax.fill(angles, percentages_plot, alpha=0.25, color='#1976d2')
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels([name.replace(' ', '\n') for name in criteria_names], fontsize=10)
        ax.set_ylim(0, 100)
        ax.set_yticks([20, 40, 60, 80, 100])
        ax.set_yticklabels(['20%', '40%', '60%', '80%', '100%'], fontsize=9)
        ax.set_title("AI Fit Profile", size=16, fontweight='bold', pad=30)
        ax.grid(True, alpha=0.7)
        
        # Add colored rings for score zones
        theta = np.linspace(0, 2*np.pi, 100)
        ax.fill_between(theta, 0, 40, alpha=0.1, color='red', label='Poor Zone')
        ax.fill_between(theta, 40, 60, alpha=0.1, color='orange', label='Moderate Zone')
        ax.fill_between(theta, 60, 75, alpha=0.1, color='yellow', label='Good Zone')
        ax.fill_between(theta, 75, 100, alpha=0.1, color='green', label='Excellent Zone')
        
        plt.tight_layout()
        st.pyplot(fig)

def render_export_section():
    """Render export and download functionality"""
    st.subheader("📤 Export Results")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📥 Download Detailed Report (CSV)", key="download_csv"):
            csv_data = generate_csv_report()
            st.download_button(
                label="📥 Download CSV Report",
                data=csv_data,
                file_name=f"ai_fit_assessment_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )
    
    with col2:
        if st.button("📊 Download Analysis Summary (JSON)", key="download_json"):
            json_data = generate_json_report()
            st.download_button(
                label="📥 Download JSON Report",
                data=json_data,
                file_name=f"ai_fit_analysis_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json"
            )

def generate_csv_report() -> str:
    """Generate comprehensive CSV report"""
    report_data = []
    
    # Add overall summary
    overall_score = calculate_overall_score(st.session_state.scores)
    report_data.append({
        'Section': 'OVERALL_SUMMARY',
        'Criteria': 'Total Score',
        'Score (1-5)': 'N/A',
        'Percentage': f"{overall_score:.1f}%",
        'Weight': '100%',
        'Weighted Score': f"{overall_score:.1f}",
        'Reasoning': f"Overall AI fit assessment score based on weighted criteria",
        'Recommendation': get_recommendation(overall_score),
        'Project Name': st.session_state.get('current_project_name', 'Unnamed Project'),
        'Analysis Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    })
    
    # Add individual criteria scores
    for criteria_name, criteria_info in CRITERIA.items():
        if criteria_name in st.session_state.scores:
            score = st.session_state.scores[criteria_name]
            percentage = (score - 1) / 4 * 100
            weighted_score = percentage * criteria_info['weight'] / 100
            
            report_data.append({
                'Section': 'CRITERIA_SCORES',
                'Criteria': criteria_name,
                'Score (1-5)': score,
                'Percentage': f"{percentage:.1f}%",
                'Weight': f"{criteria_info['weight']}%",
                'Weighted Score': f"{weighted_score:.1f}",
                'Reasoning': st.session_state.reasoning.get(criteria_name, ""),
                'Description': criteria_info['description'],
                'Key Factors': "; ".join(criteria_info['key_factors']),
                'Analysis Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
    
    # Add NLP analysis if available
    if st.session_state.nlp_analysis:
        nlp_data = st.session_state.nlp_analysis
        
        # Text statistics
        text_stats = nlp_data.get('text_stats', {})
        report_data.append({
            'Section': 'NLP_ANALYSIS',
            'Criteria': 'Text Statistics',
            'Score (1-5)': 'N/A',
            'Percentage': 'N/A',
            'Weight': 'N/A',
            'Weighted Score': 'N/A',
            'Reasoning': f"Word Count: {text_stats.get('word_count', 0)}; Character Count: {text_stats.get('char_count', 0)}; Avg Sentence Length: {text_stats.get('avg_sentence_length', 0):.1f}",
            'Analysis Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
        
        # AI Readiness
        ai_readiness = nlp_data.get('ai_readiness', {})
        if ai_readiness:
            report_data.append({
                'Section': 'NLP_ANALYSIS',
                'Criteria': 'AI Readiness Assessment',
                'Score (1-5)': 'N/A',
                'Percentage': f"{ai_readiness.get('score', 0) * 100:.1f}%",
                'Weight': 'N/A',
                'Weighted Score': 'N/A',
                'Reasoning': f"High-value indicators: {ai_readiness.get('high_value_indicators', 0)}; Moderate indicators: {ai_readiness.get('moderate_value_indicators', 0)}; Low-value indicators: {ai_readiness.get('low_value_indicators', 0)}",
                'Analysis Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
        
        # Sentiment Analysis
        sentiment = nlp_data.get('sentiment', {})
        if sentiment:
            report_data.append({
                'Section': 'NLP_ANALYSIS',
                'Criteria': 'Sentiment Analysis',
                'Score (1-5)': 'N/A',
                'Percentage': 'N/A',
                'Weight': 'N/A',
                'Weighted Score': 'N/A',
                'Reasoning': f"Label: {sentiment.get('label', 'Unknown')}; Confidence: {sentiment.get('confidence', 0):.2f}; Polarity: {sentiment.get('polarity', 0):.2f}; Consistency: {sentiment.get('consistency', 0):.2f}",
                'Analysis Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
        
        # Top Keywords
        keywords = nlp_data.get('keywords', [])
        if keywords:
            top_keywords = [f"{kw[0]} ({kw[1]:.3f})" for kw in keywords[:10]]
            report_data.append({
                'Section': 'NLP_ANALYSIS',
                'Criteria': 'Top Keywords (TF-IDF)',
                'Score (1-5)': 'N/A',
                'Percentage': 'N/A',
                'Weight': 'N/A',
                'Weighted Score': 'N/A',
                'Reasoning': "; ".join(top_keywords),
                'Analysis Date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
    
    # Convert to CSV
    df = pd.DataFrame(report_data)
    return df.to_csv(index=False)

def generate_json_report() -> str:
    """Generate comprehensive JSON report"""
    overall_score = calculate_overall_score(st.session_state.scores)
    
    report = {
        'metadata': {
            'project_name': st.session_state.get('current_project_name', 'Unnamed Project'),
            'analysis_date': datetime.now().isoformat(),
            'tool_version': '4.0',
            'analysis_type': 'comprehensive_ai_fit_assessment'
        },
        'overall_assessment': {
            'score': overall_score,
            'recommendation': get_recommendation(overall_score),
            'confidence_level': get_confidence_level(overall_score)
        },
        'criteria_scores': {},
        'nlp_analysis': st.session_state.nlp_analysis,
        'detailed_reasoning': st.session_state.reasoning
    }
    
    # Add criteria scores
    for criteria_name, criteria_info in CRITERIA.items():
        if criteria_name in st.session_state.scores:
            score = st.session_state.scores[criteria_name]
            percentage = (score - 1) / 4 * 100
            weighted_score = percentage * criteria_info['weight'] / 100
            
            report['criteria_scores'][criteria_name] = {
                'raw_score': score,
                'percentage': percentage,
                'weight': criteria_info['weight'],
                'weighted_score': weighted_score,
                'description': criteria_info['description'],
                'reasoning': st.session_state.reasoning.get(criteria_name, "")
            }
    
    return json.dumps(report, indent=2, default=str)

def get_recommendation(score: float) -> str:
    """Get recommendation based on score"""
    if score >= 70:
        return "Highly Recommended - Strong AI fit with excellent potential"
    elif score >= 60:
        return "Recommended - Good AI fit with proper planning"
    elif score >= 40:
        return "Proceed with Caution - Moderate fit, evaluate carefully"
    else:
        return "Not Recommended - Consider alternative approaches"

def get_confidence_level(score: float) -> str:
    """Get confidence level based on score distribution"""
    if not st.session_state.scores:
        return "Low"
    
    scores = list(st.session_state.scores.values())
    score_std = np.std(scores)
    
    if score_std < 0.5:
        return "High"
    elif score_std < 1.0:
        return "Medium"
    else:
        return "Low"

# Run the main application
if __name__ == "__main__":
    main()
    
    # Display results if any scores exist
    if st.session_state.scores:
        display_results()